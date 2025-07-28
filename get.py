from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import keyboard

# 配置 Selenium，启动浏览器
options = Options()
options.add_argument('--headless')  # 无头模式，不打开窗口
options.add_argument('--disable-gpu')
driver = webdriver.Chrome(options=options)  # 确保 chromedriver 在 PATH 中

# 飞书开发平台的 API 概览页面
start_url = "https://open.feishu.cn/document/client-docs/docs-add-on/05-api-doc/05-api-doc"

driver.get(start_url)
time.sleep(5)  # 等页面加载

# 滚动页面确保目录加载完整
driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
time.sleep(2)

# 找到目录区中所有链接（云文档小组件 API下的）
api_links = []
elements = driver.find_elements(By.CSS_SELECTOR, 'a[href^="/document/client-docs/docs-add-on/05-api-doc/"]')

for el in elements:
    href = el.get_attribute('href')
    if href and href not in api_links:
        api_links.append(href)

print(f"共发现 {len(api_links)} 个 API 子页面")

def create_document_with_styles():
    """创建带有预定义样式的文档"""
    doc = Document()
    
    # 添加代码块样式
    style = doc.styles.add_style('Code', 1)
    font = style.font
    font.name = 'Courier New'
    font.size = Pt(9)
    
    return doc

def save_current_page(driver, doc):
    """保存当前页面内容到Word文档"""
    try:
        wait = WebDriverWait(driver, 10)
        
        # 获取当前URL
        current_url = driver.current_url
        print(f"\n当前页面URL: {current_url}")
        
        # 等待页面加载完成
        time.sleep(3)  # 给页面更多加载时间
        
        # 获取页面标题
        try:
            title = driver.title
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            print(f"页面标题: {title}")
            
            # 添加URL作为引用
            url_paragraph = doc.add_paragraph(f"页面地址: {current_url}")
            url_paragraph.style = doc.styles['Quote']
            doc.add_paragraph()  # 添加空行
            
        except:
            doc.add_heading("未知标题", level=1)
        
        # 获取主要内容
        try:
            # 尝试获取文档内容
            # 首先尝试获取主要内容区域
            content = None
            possible_content_selectors = [
                'article',  # 通用文章容器
                '.doc-article-inner',  # 飞书特定类
                '.main-content',  # 主要内容区域
                '#docContent',  # 文档内容ID
                '.article-content',  # 文章内容类
                '.markdown-body'  # Markdown内容
            ]
            
            for selector in possible_content_selectors:
                try:
                    content = driver.find_element(By.CSS_SELECTOR, selector)
                    print(f"找到内容区域: {selector}")
                    break
                except:
                    continue
            
            if content is None:
                content = driver.find_element(By.TAG_NAME, 'body')
            
            # 获取所有标题
            for level in range(1, 7):
                headers = content.find_elements(By.TAG_NAME, f'h{level}')
                for header in headers:
                    text = header.text.strip()
                    if text and text != title:  # 避免重复添加页面标题
                        doc.add_heading(text, level=level)
            
            # 获取所有文本内容
            paragraphs = content.find_elements(By.TAG_NAME, 'p')
            for p in paragraphs:
                text = p.text.strip()
                if text:
                    doc.add_paragraph(text)
            
            # 获取所有代码块
            code_selectors = ['.code-block', 'pre', 'code', '.highlight']
            for selector in code_selectors:
                code_blocks = content.find_elements(By.CSS_SELECTOR, selector)
                for code in code_blocks:
                    text = code.text.strip()
                    if text:
                        p = doc.add_paragraph()
                        p.style = doc.styles['Code']
                        p.add_run(text)
            
            # 获取所有表格
            tables = content.find_elements(By.TAG_NAME, 'table')
            for table in tables:
                rows = table.find_elements(By.TAG_NAME, 'tr')
                if rows:
                    # 获取最大列数
                    max_cols = max(len(row.find_elements(By.TAG_NAME, 'td') or row.find_elements(By.TAG_NAME, 'th')) for row in rows)
                    if max_cols > 0:
                        doc_table = doc.add_table(rows=len(rows), cols=max_cols)
                        doc_table.style = 'Table Grid'  # 添加表格边框
                        for i, row in enumerate(rows):
                            cells = row.find_elements(By.TAG_NAME, 'td') or row.find_elements(By.TAG_NAME, 'th')
                            for j, cell in enumerate(cells):
                                doc_table.cell(i, j).text = cell.text.strip()
                        doc.add_paragraph()  # 在表格后添加空行
            
            # 获取所有列表
            lists = content.find_elements(By.CSS_SELECTOR, 'ul, ol')
            for lst in lists:
                items = lst.find_elements(By.TAG_NAME, 'li')
                for item in items:
                    text = item.text.strip()
                    if text:
                        style = 'List Bullet' if lst.tag_name == 'ul' else 'List Number'
                        doc.add_paragraph(text, style=style)
            
            print("✅ 已提取结构化内容")
            
        except Exception as e:
            print(f"提取结构化内容时出错: {str(e)}")
            print("尝试获取页面所有文本...")
            try:
                # 如果无法获取结构化内容，至少保存页面文本
                main_text = driver.find_element(By.TAG_NAME, 'body').text
                # 按行分割文本，并移除重复的空行
                lines = [line.strip() for line in main_text.split('\n')]
                current_empty_lines = 0
                for line in lines:
                    if line:
                        if current_empty_lines > 0:
                            doc.add_paragraph()  # 最多添加一个空行
                            current_empty_lines = 0
                        doc.add_paragraph(line)
                    else:
                        current_empty_lines += 1
                print("✅ 已保存页面文本")
            except Exception as e:
                print(f"保存文本时出错: {str(e)}")
        
        # 添加分页符
        doc.add_page_break()
        print("✅ 页面内容已保存到文档")
        
    except Exception as e:
        print(f"保存页面时出错: {str(e)}")

def get_feishu_doc():
    print("\n=== 飞书文档保存工具 ===")
    print("\n使用说明：")
    print("1. 程序会打开Chrome浏览器")
    print("2. 您可以在浏览器中自由导航到任何需要保存的飞书文档页面")
    print("3. 确保页面完全加载后，在终端中：")
    print("   - 按回车键 → 保存当前页面")
    print("   - 输入 's' → 显示当前状态")
    print("   - 输入 'q' → 退出程序并保存文档")
    print("\n提示：")
    print("- 建议等待页面完全加载后再保存")
    print("- 如果保存失败，可以尝试刷新页面后重新保存")
    print("- 可以保存多个页面，每个页面都会自动添加到文档中")
    print("- 如果页面内容不完整，可以尝试再次保存该页面")
    print("\n准备开始...")
    
    # 创建带有样式的Word文档
    doc = create_document_with_styles()
    doc.add_heading('飞书开放平台文档', 0)
    
    # 创建 Chrome 选项
    options = webdriver.ChromeOptions()
    
    try:
        # 创建 Chrome 浏览器实例
        driver = webdriver.Chrome(options=options)
        
        # 访问飞书文档页面
        url = "https://open.feishu.cn/document/client-docs/docs-add-on/05-api-doc/05-api-doc"
        driver.get(url)
        print("\n✅ 浏览器已启动，请导航到需要保存的页面")
        
        # 保存的页面计数
        page_count = 0
        saved_urls = set()  # 用于跟踪已保存的页面
        
        # 等待用户操作
        while True:
            user_input = input("\n请输入命令 (回车:保存当前页面, s:显示状态, q:退出): ").lower()
            
            if user_input == 'q':
                break
            elif user_input == 's':
                print(f"\n当前状态:")
                print(f"- 已保存页面数: {page_count}")
                print(f"- 当前页面URL: {driver.current_url}")
                print(f"- 当前页面标题: {driver.title}")
                if saved_urls:
                    print("\n已保存的页面:")
                    for url in sorted(saved_urls):
                        print(f"- {url}")
            else:
                current_url = driver.current_url
                if current_url in saved_urls:
                    print(f"\n警告：当前页面已经保存过。")
                    confirm = input("是否要再次保存？(y/n): ").lower()
                    if confirm != 'y':
                        continue
                
                save_current_page(driver, doc)
                page_count += 1
                saved_urls.add(current_url)
        
        # 保存文档
        output_path = 'feishu_card.docx'
        doc.save(output_path)
        print(f"\n✅ 文档已保存为 {output_path}")
        print(f"- 共保存了 {page_count} 个页面")
        
        # 关闭浏览器
        driver.quit()
        
    except Exception as e:
        print(f"发生错误: {str(e)}")
        try:
            driver.quit()
        except:
            pass

if __name__ == "__main__":
    get_feishu_doc()